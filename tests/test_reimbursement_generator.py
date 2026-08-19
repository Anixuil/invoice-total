from decimal import Decimal
from pathlib import Path
import tempfile
import unittest

from docx import Document
import fitz

from reimbursement_generator import (
    amount_to_chinese,
    amount_to_template_digits,
    parse_reimbursement_docx,
    parse_reimbursement_docx_many,
    render_reimbursement_pdf,
    validate_reimbursement_pdf,
)


class ReimbursementGeneratorTests(unittest.TestCase):
    def test_amount_to_chinese(self):
        self.assertEqual(amount_to_chinese(Decimal("321.15")), "叁佰贰拾壹元壹角伍分")
        self.assertEqual(amount_to_chinese(Decimal("100.00")), "壹佰元整")

    def test_amount_to_template_digits(self):
        self.assertEqual(
            amount_to_template_digits(Decimal("321.15")),
            ["×", "×", "×", "叁", "贰", "壹", "壹", "伍"],
        )

    def test_contract_number_label_accepts_spacing_and_variants(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "contract.docx"
            document = Document()
            for value in ("合同号 / 立项号：", "PRDP2024063"):
                document.add_paragraph(value)
            document.save(source)
            reimbursement = parse_reimbursement_docx(source)
            self.assertEqual(reimbursement.fields["contract_number"], "PRDP2024063")

    def test_contract_number_after_empty_customer_field_is_not_skipped(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "consecutive-labels.docx"
            document = Document()
            for value in ("所属客户:", "合同号/立项号:", "PRDP2024063"):
                document.add_paragraph(value)
            document.save(source)
            reimbursement = parse_reimbursement_docx(source)
            self.assertEqual(reimbursement.fields["client"], "")
            self.assertEqual(reimbursement.fields["contract_number"], "PRDP2024063")
        self.assertEqual(
            amount_to_template_digits(Decimal("101.00")),
            ["×", "×", "×", "壹", "零", "壹", "零", "零"],
        )

    def test_parses_multiple_reimbursements_in_one_docx(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "multiple.docx"
            document = Document()
            for number, claimant, amount in (("TEST-001", "张三", "12.50"), ("TEST-002", "李四", "20.00")):
                for value in (
                    "报销编号:", number, "报销人:", claimant, "所属部门:", "研发部",
                    "报销总金额:", amount, "报销明细", "报销明细1", "类型", "交通费",
                    "用途:", "出行", "金额:", amount,
                ):
                    document.add_paragraph(value)
            document.save(source)

            reimbursements = parse_reimbursement_docx_many(source)

            self.assertEqual(len(reimbursements), 2)
            self.assertEqual(reimbursements[0].fields["reimbursement_number"], "TEST-001")
            self.assertEqual(reimbursements[0].fields["claimant"], "张三")
            self.assertEqual(reimbursements[0].total, Decimal("12.50"))
            self.assertEqual(reimbursements[1].fields["reimbursement_number"], "TEST-002")
            self.assertEqual(reimbursements[1].fields["claimant"], "李四")
            self.assertEqual(reimbursements[1].total, Decimal("20.00"))

    def test_parses_docx_and_renders_multiple_pages(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "source.docx"
            document = Document()
            for value in ("报销人:", "测试人员", "所属部门:", "研发部", "报销编号:", "TEST-001", "报销总金额:", "50.00", "报销明细"):
                document.add_paragraph(value)
            for index in range(5):
                for value in (f"报销明细{index + 1}", "类型", "交通费", "用途:", f"出行{index + 1}", "金额:", "10.00"):
                    document.add_paragraph(value)
            document.save(source)
            reimbursement = parse_reimbursement_docx(source)
            self.assertEqual(len(reimbursement.details), 5)
            self.assertEqual(reimbursement.total, Decimal("50.00"))
            output = Path(directory) / "output.pdf"
            generated_at = __import__("datetime").datetime(2026, 8, 14)
            render_reimbursement_pdf(reimbursement, output, generated_at=generated_at)
            validation = validate_reimbursement_pdf(reimbursement, output, generated_at)
            self.assertTrue(validation["ok"], validation["errors"])
            with fitz.open(output) as pdf:
                self.assertEqual(pdf.page_count, 2)
                self.assertAlmostEqual(pdf[0].rect.width, 728.52, places=2)
                self.assertAlmostEqual(pdf[0].rect.height, 515.88, places=2)
                self.assertIn("测试人员", pdf[0].get_text())


if __name__ == "__main__":
    unittest.main()
